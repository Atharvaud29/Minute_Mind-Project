from docx import Document
from docx.shared import Pt
import os

FILES_DIR = os.path.join(os.getcwd(), "files")
os.makedirs(FILES_DIR, exist_ok=True)

def generate_mom_document(meeting_data, tasks, conflicts, transcript_segments=None):
    """
    Generate MoM DOCX with:
    - Meeting Details
    - Summary
    - Action Items (Tasks)
    - Conflicts (if any)
    """

    doc = Document()

    # Main Title
    title = doc.add_heading("Minutes of Meeting (MoM)", level=1)
    title.alignment = 1

    # -----------------------------
    # MEETING DETAILS
    # -----------------------------
    doc.add_paragraph(f"Meeting Name: {meeting_data.title}")
    doc.add_paragraph(f"Meeting Host: {meeting_data.host}")

    doc.add_paragraph("Present Members:")
    if meeting_data.presentees:
        for p in meeting_data.presentees.split(","):
            doc.add_paragraph(f"    • {p.strip()}")
    else:
        doc.add_paragraph("    • Not Provided")

    doc.add_paragraph("\nAbsent Members:")
    if meeting_data.absentees:
        for p in meeting_data.absentees.split(","):
            doc.add_paragraph(f"    • {p.strip()}")
    else:
        doc.add_paragraph("    • None")

    doc.add_paragraph(f"\nDate of Meeting: {meeting_data.date}")
    doc.add_paragraph(f"Time: {meeting_data.start_time or ''} – {meeting_data.end_time or ''} (IST)")

    # -----------------------------
    # SUMMARY SECTION
    # -----------------------------
    doc.add_heading("Agenda → Summary of Meeting :", level=2)

    if meeting_data.summary:
        # Split into bullet points if possible
        points = [
            s.strip() for s in meeting_data.summary.replace("\n", " ").split(".")
            if s.strip()
        ]
        for p in points:
            doc.add_paragraph(f"    • {p}")
    else:
        doc.add_paragraph("    • No summary available.")

    # -----------------------------
    # TASKS SECTION
    # -----------------------------
    doc.add_heading("Action Items → Tasks Assigned :", level=2)

    if tasks:
        for t in tasks:
            person = t.person or "Unknown"
            task_text = t.task or "Unnamed Task"
            status = t.status or "Pending"
            doc.add_paragraph(f"    • {person} → {task_text}. ({status})")
    else:
        doc.add_paragraph("    • No tasks assigned.")

    # -----------------------------
    # CONFLICTS SECTION
    # -----------------------------
    doc.add_heading("Conflicts / Issues Raised :", level=2)

    if conflicts:
        for c in conflicts:
            issue = c.issue or "No issue description"
            raised_by = c.raised_by or "Unknown"
            sev = c.severity
            doc.add_paragraph(f"    • {issue} → Raised By: {raised_by}  (Severity: {sev})")
    else:
        doc.add_paragraph("    • No conflicts reported.")

    # -----------------------------
    # SAVE FILE
    # -----------------------------
    filename = f"MoM_{meeting_data.id}.docx"
    filepath = os.path.join(FILES_DIR, filename)
    doc.save(filepath)

    return filepath